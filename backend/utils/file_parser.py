"""
파일 형식별 텍스트 추출 유틸리티
지원 형식: PDF, DOCX, XLSX, TXT
"""

import io
import logging
from pathlib import Path
import re
from typing import Optional, Union, List, Dict, Any, Tuple

import PyPDF2
import pdfplumber
from docx import Document
from openpyxl import load_workbook
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.merge import MergedCellRange

logger = logging.getLogger(__name__)


class MergedCellProcessor:
    """병합된 셀 처리 유틸리티 클래스"""
    
    @staticmethod
    def unmerge_cells_and_fill_values(sheet) -> Dict[Tuple[int, int], Any]:
        """
        병합된 셀을 해제하고 모든 셀에 값을 채움
        
        Args:
            sheet: openpyxl 워크시트 객체
            
        Returns:
            Dict[Tuple[int, int], Any]: (row, col) -> value 매핑
        """
        logger.info("병합된 셀 처리 시작...")
        
        # 모든 셀의 값을 저장할 딕셔너리
        cell_values = {}
        
        # 먼저 모든 셀의 값을 읽어옴
        for row in sheet.iter_rows():
            for cell in row:
                if cell.value is not None:
                    cell_values[(cell.row, cell.column)] = cell.value
        
        # 병합된 셀 범위들을 처리
        merged_ranges = list(sheet.merged_cells.ranges)
        logger.info(f"병합된 셀 범위 {len(merged_ranges)}개 발견")
        
        for merged_range in merged_ranges:
            # 병합된 범위의 첫 번째 셀(왼쪽 위)에서 값 가져오기
            top_left_cell = sheet.cell(merged_range.min_row, merged_range.min_col)
            merged_value = top_left_cell.value
            
            if merged_value is not None:
                logger.debug(f"병합 범위 {merged_range} 처리: '{merged_value}'")
                
                # 병합된 범위의 모든 셀에 동일한 값 채우기
                for row in range(merged_range.min_row, merged_range.max_row + 1):
                    for col in range(merged_range.min_col, merged_range.max_col + 1):
                        cell_values[(row, col)] = merged_value
                        logger.debug(f"  셀 ({row}, {col})에 값 '{merged_value}' 복제")
        
        logger.info(f"병합된 셀 처리 완료: {len(cell_values)}개 셀 값 설정")
        return cell_values
    
    @staticmethod
    def process_table_with_merged_cells(table_data: List[List[str]]) -> List[List[str]]:
        """
        PDF 테이블 데이터에서 병합된 셀 패턴을 감지하고 값을 복제
        
        Args:
            table_data: PDF에서 추출한 테이블 데이터
            
        Returns:
            List[List[str]]: 병합된 셀이 해제된 테이블 데이터
        """
        if not table_data or len(table_data) < 2:
            return table_data
        
        logger.info("PDF 테이블 병합된 셀 처리 시작...")
        
        # 헤더와 데이터 분리
        headers = table_data[0]
        data_rows = table_data[1:]
        
        # 각 컬럼별로 병합된 셀 패턴 감지 및 처리
        processed_rows = []
        
        for row_idx, row in enumerate(data_rows):
            processed_row = []
            
            for col_idx, cell_value in enumerate(row):
                if cell_value and cell_value.strip():
                    # 값이 있는 경우 그대로 사용
                    processed_row.append(cell_value.strip())
                else:
                    # 빈 셀인 경우 이전 행에서 같은 컬럼의 값을 찾아서 복제
                    filled_value = MergedCellProcessor._find_previous_value(
                        data_rows, row_idx, col_idx
                    )
                    processed_row.append(filled_value)
            
            processed_rows.append(processed_row)
        
        # 헤더와 처리된 데이터 결합
        result = [headers] + processed_rows
        
        logger.info(f"PDF 테이블 병합된 셀 처리 완료: {len(result)}개 행")
        return result
    
    @staticmethod
    def _find_previous_value(data_rows: List[List[str]], current_row_idx: int, col_idx: int) -> str:
        """
        현재 행 이전에서 같은 컬럼의 값을 찾아서 반환
        
        Args:
            data_rows: 데이터 행들
            current_row_idx: 현재 행 인덱스
            col_idx: 컬럼 인덱스
            
        Returns:
            str: 찾은 값 또는 빈 문자열
        """
        # 현재 행 이전의 행들을 역순으로 검색
        for prev_row_idx in range(current_row_idx - 1, -1, -1):
            if prev_row_idx < len(data_rows):
                prev_row = data_rows[prev_row_idx]
                if col_idx < len(prev_row) and prev_row[col_idx] and prev_row[col_idx].strip():
                    logger.debug(f"병합된 셀 값 복제: 행 {prev_row_idx + 1} → 행 {current_row_idx + 1}, 컬럼 {col_idx + 1}")
                    return prev_row[col_idx].strip()
        
        return ""


class FileParser:
    """파일 형식별 텍스트 추출기"""
    
    SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.xlsx', '.txt'}
    MAX_FILE_SIZE = 10 * 1024 * 1024  # 10MB
    
    @classmethod
    async def extract_text(cls, file_content: bytes, filename: str) -> Union[str, List[Dict[str, Any]]]:
        """
        파일 내용에서 텍스트 추출
        
        Args:
            file_content: 파일 바이트 데이터
            filename: 파일명 (확장자 포함)
            
        Returns:
            XLSX: 구조화된 셀 데이터 리스트
            기타: 추출된 텍스트 문자열
            
        Raises:
            ValueError: 지원하지 않는 파일 형식
            RuntimeError: 텍스트 추출 실패
        """
        if len(file_content) > FileParser.MAX_FILE_SIZE:
            raise ValueError(f"파일 크기가 {FileParser.MAX_FILE_SIZE / 1024 / 1024}MB를 초과합니다")
        
        file_ext = Path(filename).suffix.lower()
        
        if file_ext not in FileParser.SUPPORTED_EXTENSIONS:
            raise ValueError(f"지원하지 않는 파일 형식: {file_ext}")
        
        try:
            if file_ext == '.pdf':
                return await FileParser._extract_from_pdf(file_content)
            elif file_ext == '.docx':
                return await FileParser._extract_from_docx(file_content)
            elif file_ext == '.xlsx':
                return await FileParser._extract_from_xlsx(file_content)
            elif file_ext == '.txt':
                return await FileParser._extract_from_txt(file_content)
        except Exception as e:
            logger.error(f"텍스트 추출 실패 - 파일: {filename}, 오류: {str(e)}")
            raise RuntimeError(f"텍스트 추출 실패: {str(e)}")
    
    @staticmethod
    async def _extract_from_pdf(file_content: bytes) -> Union[str, List[Dict[str, Any]]]:
        """PDF 파일에서 텍스트 및 표 추출"""
        try:
            pdf_file = io.BytesIO(file_content)
            
            # pdfplumber로 표 감지 및 추출 시도
            try:
                with pdfplumber.open(pdf_file) as pdf:
                    structured_data = []
                    has_tables = False
                    
                    for page_num, page in enumerate(pdf.pages):
                        logger.info(f"PDF 페이지 {page_num + 1} 처리 중...")
                        
                        # 표 추출 시도
                        tables = page.extract_tables()
                        if tables:
                            logger.info(f"페이지 {page_num + 1}에서 {len(tables)}개 표 발견")
                            has_tables = True
                            
                            for table_idx, table in enumerate(tables):
                                if table and len(table) > 1:  # 헤더 + 데이터 행이 있는 경우
                                    logger.info(f"표 {table_idx + 1} 처리 중... (행 수: {len(table)})")
                                    
                                    # 병합된 셀 처리 적용
                                    processed_table = MergedCellProcessor.process_table_with_merged_cells(table)
                                    logger.info(f"표 {table_idx + 1} 병합된 셀 처리 완료")
                                    
                                    table_data = FileParser._process_pdf_table(processed_table, page_num + 1, table_idx + 1)
                                    logger.info(f"표 {table_idx + 1}에서 {len(table_data)}개 항목 추출")
                                    structured_data.extend(table_data)
                                else:
                                    logger.warning(f"표 {table_idx + 1}이 비어있거나 데이터가 부족합니다")
                        
                        # 표가 없는 경우 일반 텍스트 추출
                        if not tables:
                            logger.info(f"페이지 {page_num + 1}에 표가 없음, 텍스트 추출 시도")
                            page_text = page.extract_text()
                            if page_text and page_text.strip():
                                logger.info(f"페이지 {page_num + 1}에서 텍스트 추출 성공 ({len(page_text)}자)")
                                # 텍스트를 구조화된 형태로 변환
                                text_data = FileParser._process_pdf_text(page_text, page_num + 1)
                                logger.info(f"텍스트에서 {len(text_data)}개 항목 추출")
                                structured_data.extend(text_data)
                            else:
                                logger.warning(f"페이지 {page_num + 1}에서 텍스트를 추출할 수 없습니다")
                    
                    if has_tables and structured_data:
                        logger.info(f"PDF 표 데이터 추출 완료: {len(structured_data)}개 항목")
                        return structured_data
                    elif structured_data:
                        logger.info(f"PDF 텍스트 데이터 추출 완료: {len(structured_data)}개 항목")
                        return structured_data
                    else:
                        # 기존 방식으로 폴백
                        return await FileParser._extract_from_pdf_fallback(file_content)
                        
            except Exception as e:
                logger.warning(f"pdfplumber 처리 실패, 기존 방식으로 폴백: {str(e)}")
                return await FileParser._extract_from_pdf_fallback(file_content)
                
        except Exception as e:
            raise RuntimeError(f"PDF 처리 오류: {str(e)}")
    
    @staticmethod
    async def _extract_from_pdf_fallback(file_content: bytes) -> str:
        """PDF 파일에서 텍스트 추출 (기존 방식)"""
        text_parts = []
        
        try:
            pdf_file = io.BytesIO(file_content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            for page_num, page in enumerate(pdf_reader.pages):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_parts.append(page_text)
                except Exception as e:
                    logger.warning(f"PDF 페이지 {page_num + 1} 텍스트 추출 실패: {str(e)}")
                    continue
            
            if not text_parts:
                raise RuntimeError("PDF에서 텍스트를 추출할 수 없습니다")
                
            return '\n\n'.join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"PDF 처리 오류: {str(e)}")
    
    @staticmethod
    def _process_pdf_table(table: List[List[str]], page_num: int, table_idx: int) -> List[Dict[str, Any]]:
        """PDF 표를 구조화된 데이터로 변환"""
        if not table or len(table) < 2:
            logger.warning(f"표가 비어있거나 데이터가 부족합니다 (행 수: {len(table) if table else 0})")
            return []
        
        logger.info(f"표 처리 시작 - 페이지: {page_num}, 표: {table_idx}, 행 수: {len(table)}")
        
        structured_data = []
        headers = table[0] if table else []
        
        # 헤더 정리
        clean_headers = []
        for header in headers:
            if header:
                clean_headers.append(str(header).strip())
            else:
                clean_headers.append("")
        
        logger.info(f"표 헤더: {clean_headers}")
        
        # 데이터 행 처리
        processed_rows = 0
        for row_idx, row in enumerate(table[1:], 1):
            if not row or all(not cell for cell in row):
                logger.debug(f"행 {row_idx} 건너뛰기 (빈 행)")
                continue
            
            # 행 데이터 정리
            clean_row = []
            for cell in row:
                if cell:
                    clean_row.append(str(cell).strip())
                else:
                    clean_row.append("")
            
            # 계층형 구조 추출 시도
            lvl1, lvl2, lvl3, lvl4 = FileParser._extract_hierarchical_structure(clean_headers, clean_row)
            
            if lvl1 or lvl2 or lvl3 or lvl4:
                logger.debug(f"행 {row_idx} 계층형 구조: lvl1={lvl1}, lvl2={lvl2}, lvl3={lvl3}, lvl4={lvl4}")
            
            # 각 셀별로 데이터 생성
            row_cells = 0
            for col_idx, (header, value) in enumerate(zip(clean_headers, clean_row)):
                if value:  # 값이 있는 셀만 처리
                    cell_data = {
                        "page": page_num,
                        "table": table_idx,
                        "row": row_idx,
                        "col": col_idx + 1,
                        "header": header,
                        "value": value,
                        "lvl1": lvl1,
                        "lvl2": lvl2,
                        "lvl3": lvl3,
                        "lvl4": lvl4,
                        "row_context": " | ".join([v for v in clean_row if v])
                    }
                    structured_data.append(cell_data)
                    row_cells += 1
            
            if row_cells > 0:
                processed_rows += 1
                logger.debug(f"행 {row_idx} 처리 완료: {row_cells}개 셀")
        
        logger.info(f"표 처리 완료: {processed_rows}개 행, {len(structured_data)}개 셀")
        return structured_data
    
    @staticmethod
    def _process_pdf_text(text: str, page_num: int) -> List[Dict[str, Any]]:
        """PDF 텍스트를 구조화된 데이터로 변환"""
        if not text or not text.strip():
            return []
        
        lines = text.strip().split('\n')
        structured_data = []
        
        for line_idx, line in enumerate(lines, 1):
            if line.strip():
                # 간단한 구조화 (실제로는 더 정교한 파싱이 필요할 수 있음)
                structured_data.append({
                    "page": page_num,
                    "table": 0,
                    "row": line_idx,
                    "col": 1,
                    "header": "텍스트",
                    "value": line.strip(),
                    "lvl1": "",
                    "lvl2": "",
                    "lvl3": "",
                    "lvl4": line.strip(),
                    "row_context": line.strip()
                })
        
        return structured_data
    
    @staticmethod
    def _extract_hierarchical_structure(headers: List[str], row: List[str]) -> tuple:
        """표 행에서 계층형 구조 추출"""
        lvl1 = lvl2 = lvl3 = lvl4 = ""
        
        # 괄호 안 내용 추출 함수
        def _extract_bracket_content(text: str) -> str:
            """조항에서 괄호 안의 내용만 추출"""
            if not text:
                return ""
            
            import re
            # "제N조(내용)" 패턴에서 괄호 안의 내용만 추출
            article_match = re.match(r'제\d+조\s*\(([^)]+)\)', text)
            if article_match:
                content = article_match.group(1).strip()
                if content:  # 괄호 안에 실제 내용이 있는 경우만
                    logger.debug(f"표에서 조항 내용 추출: '{text}' → '{content}'")
                    return content
                else:
                    logger.debug(f"표에서 빈 괄호 조항 제외: {text}")
                    return ""
            
            # 일반 텍스트는 그대로 반환
            return text
        
        # 일반적인 계층형 구조 패턴 감지
        if len(headers) >= 4 and len(row) >= 4:
            # 첫 번째 컬럼이 구분/카테고리인 경우
            if headers[0] and any(keyword in headers[0].lower() for keyword in ['구분', '분류', '카테고리', '구분1']):
                lvl1 = _extract_bracket_content(row[0]) if row[0] else ""
            if len(headers) > 1 and headers[1] and any(keyword in headers[1].lower() for keyword in ['구분2', '세부', '하위']):
                lvl2 = _extract_bracket_content(row[1]) if row[1] else ""
            if len(headers) > 2 and headers[2] and any(keyword in headers[2].lower() for keyword in ['구분3', '세부', '항목']):
                lvl3 = _extract_bracket_content(row[2]) if row[2] else ""
            if len(headers) > 3 and headers[3] and any(keyword in headers[3].lower() for keyword in ['내용', '세부내용', '설명']):
                lvl4 = _extract_bracket_content(row[3]) if row[3] else ""
        
        # 자동 감지가 실패한 경우 첫 4개 컬럼을 lvl1~lvl4로 매핑
        if not lvl1 and len(row) >= 1:
            lvl1 = _extract_bracket_content(row[0]) if row[0] else ""
        if not lvl2 and len(row) >= 2:
            lvl2 = _extract_bracket_content(row[1]) if row[1] else ""
        if not lvl3 and len(row) >= 3:
            lvl3 = _extract_bracket_content(row[2]) if row[2] else ""
        if not lvl4 and len(row) >= 4:
            lvl4 = _extract_bracket_content(row[3]) if row[3] else ""
        
        return lvl1, lvl2, lvl3, lvl4
    
    @staticmethod
    async def _extract_from_docx(file_content: bytes) -> Union[str, List[Dict[str, Any]]]:
        """DOCX 파일에서 텍스트 및 구조 추출"""
        try:
            docx_file = io.BytesIO(file_content)
            doc = Document(docx_file)
            
            logger.info("DOCX 파일 구조 분석 시작")
            
            # 구조화된 데이터 추출 시도
            structured_data = []
            
            # 1. 문단 구조 분석
            paragraph_data = FileParser._extract_docx_paragraphs(doc.paragraphs)
            if paragraph_data:
                structured_data.extend(paragraph_data)
                logger.info(f"문단에서 {len(paragraph_data)}개 항목 추출")
            
            # 2. 표 구조 분석 (병합된 셀 처리 적용)
            table_data = FileParser._extract_docx_tables(doc.tables)
            if table_data:
                structured_data.extend(table_data)
                logger.info(f"표에서 {len(table_data)}개 항목 추출 (병합된 셀 처리 적용)")
            
            if structured_data:
                logger.info(f"DOCX 구조화 데이터 추출 완료: {len(structured_data)}개 항목")
                return structured_data
            else:
                # 구조화 실패 시 기존 방식으로 폴백
                logger.warning("구조화 추출 실패, 기존 텍스트 방식으로 폴백")
                return await FileParser._extract_from_docx_fallback(doc)
                
        except Exception as e:
            logger.error(f"DOCX 처리 오류: {str(e)}")
            raise RuntimeError(f"DOCX 처리 오류: {str(e)}")
    
    @staticmethod
    async def _extract_from_docx_fallback(doc) -> str:
        """DOCX 파일에서 텍스트 추출 (기존 방식)"""
        try:
            text_parts = []
            
            # 문단 텍스트 추출
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            # 표 텍스트 추출
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_parts.append(' | '.join(row_text))
            
            if not text_parts:
                raise RuntimeError("DOCX에서 텍스트를 추출할 수 없습니다")
            
            return '\n'.join(text_parts)
            
        except Exception as e:
            raise RuntimeError(f"DOCX 폴백 처리 오류: {str(e)}")
    
    @staticmethod
    def _extract_docx_paragraphs(paragraphs) -> List[Dict[str, Any]]:
        """DOCX 문단에서 구조화된 데이터 추출"""
        structured_data = []
        current_lvl1 = ""
        current_lvl2 = ""
        current_lvl3 = ""
        
        logger.info(f"문단 처리 시작: {len(paragraphs)}개 문단")
        
        for para_idx, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()
            if not text:
                continue
            
            # 제목/조항 패턴 감지
            lvl1, lvl2, lvl3, lvl4 = FileParser._extract_docx_structure(text, paragraph)
            
            # 추가 안전장치: 각 레벨에서 괄호 내용 추출
            def _extract_bracket_content_safe(text: str) -> str:
                """조항에서 괄호 안의 내용만 추출 (추가 안전장치)"""
                if not text:
                    return ""
                
                import re
                # "제N조(내용)" 패턴에서 괄호 안의 내용만 추출
                article_match = re.match(r'제\d+조\s*\(([^)]+)\)', text)
                if article_match:
                    content = article_match.group(1).strip()
                    if content:  # 괄호 안에 실제 내용이 있는 경우만
                        logger.debug(f"문단에서 조항 내용 추출: '{text}' → '{content}'")
                        return content
                    else:
                        logger.debug(f"문단에서 빈 괄호 조항 제외: {text}")
                        return ""
                
                # 일반 텍스트는 그대로 반환
                return text
            
            # 각 레벨에 괄호 내용 추출 적용
            lvl1 = _extract_bracket_content_safe(lvl1)
            lvl2 = _extract_bracket_content_safe(lvl2)
            lvl3 = _extract_bracket_content_safe(lvl3)
            lvl4 = _extract_bracket_content_safe(lvl4)
            
            # 계층형 구조 업데이트 (forward fill)
            if lvl1:
                current_lvl1 = lvl1
                current_lvl2 = ""
                current_lvl3 = ""
                logger.info(f"조항 발견: {lvl1}")
            elif lvl2:
                current_lvl2 = lvl2
                current_lvl3 = ""
                logger.debug(f"소항목 발견: {lvl2}")
            elif lvl3:
                current_lvl3 = lvl3
                logger.debug(f"세부항목 발견: {lvl3}")
            
            # lvl4가 없으면 전체 텍스트를 lvl4로 사용
            if not lvl4:
                lvl4 = text
            
            # 구조화된 데이터 생성
            if text:  # 빈 텍스트가 아닌 경우만
                structured_data.append({
                    "page": 1,  # DOCX는 페이지 정보가 없으므로 1로 설정
                    "table": 0,
                    "row": para_idx + 1,
                    "col": 1,
                    "header": "문단",
                    "value": text,
                    "lvl1": current_lvl1,
                    "lvl2": current_lvl2,
                    "lvl3": current_lvl3,
                    "lvl4": lvl4,
                    "row_context": text,
                    "paragraph_style": paragraph.style.name if paragraph.style else "Normal"
                })
        
        logger.info(f"문단 처리 완료: {len(structured_data)}개 항목 추출")
        return structured_data
    
    @staticmethod
    def _extract_docx_tables(tables) -> List[Dict[str, Any]]:
        """DOCX 표에서 구조화된 데이터 추출 (병합된 셀 처리 적용)"""
        structured_data = []
        
        logger.info(f"표 처리 시작: {len(tables)}개 표")
        
        for table_idx, table in enumerate(tables, 1):
            logger.info(f"표 {table_idx} 처리 중... (행 수: {len(table.rows)})")
            
            if not table.rows:
                logger.warning(f"표 {table_idx}이 비어있습니다")
                continue
            
            # 표 데이터를 2차원 리스트로 변환
            table_data = []
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data.append(row_data)
            
            # 병합된 셀 처리 적용
            processed_table_data = MergedCellProcessor.process_table_with_merged_cells(table_data)
            logger.info(f"표 {table_idx} 병합된 셀 처리 완료")
            
            # 헤더 추출 (첫 번째 행)
            headers = processed_table_data[0] if processed_table_data else []
            
            logger.info(f"표 {table_idx} 헤더: {headers}")
            
            # 데이터 행 처리
            processed_rows = 0
            for row_idx, row_data in enumerate(processed_table_data[1:], 1):
                if not row_data:
                    continue
                
                # 계층형 구조 추출
                lvl1, lvl2, lvl3, lvl4 = FileParser._extract_hierarchical_structure(headers, row_data)
                
                if lvl1 or lvl2 or lvl3 or lvl4:
                    logger.debug(f"표 {table_idx} 행 {row_idx} 계층형 구조: lvl1={lvl1}, lvl2={lvl2}, lvl3={lvl3}, lvl4={lvl4}")
                
                # 각 셀별로 데이터 생성
                row_cells = 0
                for col_idx, (header, value) in enumerate(zip(headers, row_data)):
                    if value:  # 값이 있는 셀만 처리
                        structured_data.append({
                            "page": 1,
                            "table": table_idx,
                            "row": row_idx,
                            "col": col_idx + 1,
                            "header": header,
                            "value": value,
                            "lvl1": lvl1,
                            "lvl2": lvl2,
                            "lvl3": lvl3,
                            "lvl4": lvl4,
                            "row_context": " | ".join([v for v in row_data if v])
                        })
                        row_cells += 1
                
                if row_cells > 0:
                    processed_rows += 1
                    logger.debug(f"표 {table_idx} 행 {row_idx} 처리 완료: {row_cells}개 셀")
            
            logger.info(f"표 {table_idx} 처리 완료: {processed_rows}개 행, {len([item for item in structured_data if item.get('table') == table_idx])}개 셀")
        
        logger.info(f"표 처리 완료: {len(structured_data)}개 항목 추출 (병합된 셀 처리 적용)")
        return structured_data
    
    @staticmethod
    def _extract_docx_structure(text: str, paragraph) -> tuple:
        """DOCX 문단에서 계층형 구조 추출"""
        lvl1 = lvl2 = lvl3 = lvl4 = ""
        
        # 제목/조항 패턴 감지
        import re
        
        # 제1조, 제2조 등의 패턴 (lvl1) - 괄호 안 내용만 추출
        article_match = re.match(r'제(\d+)조\s*\(([^)]+)\)', text)
        if article_match:
            content = article_match.group(2).strip()
            if content:  # 괄호 안에 실제 내용이 있는 경우만
                lvl1 = content  # 괄호 안의 내용만 저장
                logger.debug(f"조항 내용 추출: '{text}' → '{content}'")
            else:
                logger.debug(f"빈 괄호 조항 제외: {text}")
            return lvl1, lvl2, lvl3, lvl4
        
        # ①, ②, ③ 등의 패턴 (lvl2)
        if re.match(r'[①②③④⑤⑥⑦⑧⑨⑩]', text):
            lvl2 = text
            return lvl1, lvl2, lvl3, lvl4
        
        # 1), 2), 3) 등의 패턴 (lvl3)
        if re.match(r'\d+\)', text):
            lvl3 = text
            return lvl1, lvl2, lvl3, lvl4
        
        # 가), 나), 다) 등의 패턴 (lvl3)
        if re.match(r'[가-힣]\)', text):
            lvl3 = text
            return lvl1, lvl2, lvl3, lvl4
        
        # 일반 문단은 lvl4로 처리
        lvl4 = text
        return lvl1, lvl2, lvl3, lvl4
    
    @staticmethod
    def _has_cell_border(cell) -> bool:
        """셀에 테두리가 있는지 확인"""
        if not cell or not hasattr(cell, 'border'):
            return False
        
        border = cell.border
        if not border:
            return False
        
        # 상하좌우 중 하나라도 테두리가 있으면 True
        return bool(
            border.top and border.top.style is not None or
            border.bottom and border.bottom.style is not None or
            border.left and border.left.style is not None or
            border.right and border.right.style is not None
        )
    
    @staticmethod
    def _find_table_range_by_border(sheet) -> tuple:
        """테두리를 기준으로 표 영역의 시작과 끝 행을 찾음"""
        if sheet.max_row == 0:
            return -1, -1
        
        table_start = -1
        table_end = -1
        
        # 모든 행을 순회하며 테두리가 있는 행 찾기
        for row_idx in range(1, sheet.max_row + 1):
            has_border = False
            
            # 이 행에 테두리가 있는 셀이 있는지 확인
            for col_idx in range(1, sheet.max_column + 1):
                cell = sheet.cell(row=row_idx, column=col_idx)
                if FileParser._has_cell_border(cell):
                    has_border = True
                    break
            
            if has_border:
                if table_start == -1:
                    table_start = row_idx
                table_end = row_idx
            elif table_start != -1:
                # 테두리가 있는 행이 시작되었는데 현재 행에 테두리가 없으면 표 종료
                break
        
        return table_start, table_end
    
    @staticmethod
    async def _extract_from_xlsx(file_content: bytes) -> List[Dict[str, Any]]:
        """XLSX 파일에서 테두리 기반 표 영역만 추출 (병합된 셀 처리 및 계층형 컬럼 지원)"""
        try:
            xlsx_file = io.BytesIO(file_content)
            # cellStyles=True 옵션으로 스타일 정보 포함하여 로드
            workbook = load_workbook(xlsx_file, read_only=False, data_only=True)
            
            cell_data = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                
                # 시트가 비어있으면 건너뛰기
                if sheet.max_row == 0:
                    continue
                
                logger.info(f"시트 '{sheet_name}' 처리 중... (전체 행: {sheet.max_row})")
                
                # 병합된 셀 처리 적용
                merged_cell_values = MergedCellProcessor.unmerge_cells_and_fill_values(sheet)
                logger.info(f"시트 '{sheet_name}' 병합된 셀 처리 완료")
                
                # 테두리 기반으로 표 영역 찾기
                table_start, table_end = FileParser._find_table_range_by_border(sheet)
                
                if table_start == -1 or table_end == -1:
                    logger.warning(f"시트 '{sheet_name}'에서 테두리가 있는 표를 찾을 수 없습니다")
                    continue
                
                logger.info(f"시트 '{sheet_name}'에서 표 영역 발견: 행 {table_start}~{table_end}")
                
                # 표 영역 내에서 헤더 추출 (표의 첫 번째 행)
                headers = {}
                for col_idx in range(1, sheet.max_column + 1):
                    cell_value = merged_cell_values.get((table_start, col_idx))
                    if cell_value is not None:
                        headers[col_idx] = str(cell_value).strip()
                    else:
                        headers[col_idx] = f"Column{col_idx}"
                
                logger.info(f"표 헤더: {headers}")
                
                # 고정 매핑 규칙 (항상 동일한 컬럼 순서 사용)
                # 컬럼 순서: 구분1(1번째), 구분2(2번째), 구분3(3번째), 세부 내용(4번째), 비고(5번째)
                lvl1_col_idx = 2  # 구분1 - 항상 1번째 컬럼
                lvl2_col_idx = 3  # 구분2 - 항상 2번째 컬럼
                lvl3_col_idx = 4  # 구분3 - 항상 3번째 컬럼
                detail_col_idx = 5  # 세부 내용 - 항상 4번째 컬럼
                remarks_col_idx = 6  # 비고 - 항상 5번째 컬럼
                
                # 계층형 컬럼 추출을 위한 상태 관리
                lvl1_value = None
                lvl2_value = None
                lvl3_value = None
                
                # 표 영역 내에서만 데이터 처리 (table_start+1부터 table_end까지)
                for row_idx in range(table_start + 1, table_end + 1):
                    # 병합된 셀 처리된 값들로 행 데이터 구성
                    row_values = []
                    for col_idx in range(1, sheet.max_column + 1):
                        cell_value = merged_cell_values.get((row_idx, col_idx))
                        row_values.append(cell_value)
                    
                    # 행 컨텍스트 생성 (같은 행의 모든 값들)
                    row_context_parts = []
                    for val in row_values:
                        if val is not None:
                            val_str = str(val).strip()
                            if val_str:
                                row_context_parts.append(val_str)
                    row_context = " | ".join(row_context_parts) if row_context_parts else ""
                    
                    # 계층형 컬럼 값 업데이트 (forward fill 로직)
                    def _clean_hierarchy_value(text: str) -> str:
                        """lvl1~lvl3 값에서 줄바꿈 제거 및 공백 정규화"""
                        if not text:
                            return ""
                        # 줄바꿈을 공백으로 치환 후 다중 공백 축약
                        no_newline = re.sub(r"[\r\n]+", " ", text)
                        return re.sub(r"\s+", " ", no_newline).strip()

                    current_lvl1 = lvl1_value
                    current_lvl2 = lvl2_value
                    current_lvl3 = lvl3_value
                    
                    # 구분1 컬럼 확인
                    if lvl1_col_idx and len(row_values) >= lvl1_col_idx and row_values[lvl1_col_idx-1] is not None:
                        val_str = _clean_hierarchy_value(str(row_values[lvl1_col_idx-1]).strip())
                        if val_str:
                            current_lvl1 = val_str
                            lvl1_value = val_str  # 다음 행을 위해 저장
                    
                    # 구분2 컬럼 확인
                    if lvl2_col_idx and len(row_values) >= lvl2_col_idx and row_values[lvl2_col_idx-1] is not None:
                        val_str = _clean_hierarchy_value(str(row_values[lvl2_col_idx-1]).strip())
                        if val_str:
                            current_lvl2 = val_str
                            lvl2_value = val_str
                    
                    # 구분3 컬럼 확인
                    if lvl3_col_idx and len(row_values) >= lvl3_col_idx and row_values[lvl3_col_idx-1] is not None:
                        val_str = _clean_hierarchy_value(str(row_values[lvl3_col_idx-1]).strip())
                        if val_str:
                            current_lvl3 = val_str
                            lvl3_value = val_str
                    
                    # 세부 내용과 비고 컬럼 추출
                    detail_content = ""
                    remarks = ""
                    
                    if detail_col_idx and len(row_values) >= detail_col_idx and row_values[detail_col_idx-1] is not None:
                        detail_content = str(row_values[detail_col_idx-1]).strip()
                    
                    if remarks_col_idx and len(row_values) >= remarks_col_idx and row_values[remarks_col_idx-1] is not None:
                        remarks = str(row_values[remarks_col_idx-1]).strip()
                    
                    # lvl4 생성 (세부 내용 + 비고)
                    lvl4_parts = []
                    if detail_content:
                        lvl4_parts.append(detail_content)
                    if remarks:
                        lvl4_parts.append(remarks)
                    lvl4_value = " | ".join(lvl4_parts) if lvl4_parts else ""
                    
                    # 각 셀 처리 (병합된 셀 처리된 값 사용)
                    for col_idx, cell_value in enumerate(row_values, 1):
                        if cell_value is not None:
                            cell_str = str(cell_value).strip()
                            if cell_str and len(cell_str) > 0:  # 빈 문자열이 아닌 경우
                                col_letter = get_column_letter(col_idx)
                                
                                # 숫자 여부 판단
                                is_numeric = False
                                try:
                                    float(cell_str)
                                    is_numeric = True
                                except ValueError:
                                    pass
                                
                                cell_info = {
                                    "sheet": sheet_name,
                                    "row": row_idx,
                                    "col": col_letter,
                                    "cell_address": f"{sheet_name}!{col_letter}{row_idx}",
                                    "value": cell_str,
                                    "row_context": row_context,
                                    "col_header": headers.get(col_idx, f"Column{col_idx}"),
                                    "is_numeric": is_numeric,
                                    # 계층형 컬럼 추가
                                    "lvl1": _clean_hierarchy_value(current_lvl1 or ""),
                                    "lvl2": _clean_hierarchy_value(current_lvl2 or ""),
                                    "lvl3": _clean_hierarchy_value(current_lvl3 or ""),
                                    "lvl4": lvl4_value
                                }
                                cell_data.append(cell_info)
            
            workbook.close()
            
            if not cell_data:
                raise RuntimeError("XLSX에서 데이터를 추출할 수 없습니다")
            
            logger.info(f"XLSX 셀 데이터 추출 완료: {len(cell_data)}개 셀 (병합된 셀 처리 및 계층형 컬럼 포함)")
            return cell_data
            
        except Exception as e:
            raise RuntimeError(f"XLSX 처리 오류: {str(e)}")
    
    @staticmethod
    async def _extract_from_txt(file_content: bytes) -> str:
        """TXT 파일에서 텍스트 추출"""
        try:
            # UTF-8로 먼저 시도
            try:
                text = file_content.decode('utf-8')
            except UnicodeDecodeError:
                # UTF-8 실패 시 CP949 (EUC-KR) 시도
                try:
                    text = file_content.decode('cp949')
                except UnicodeDecodeError:
                    # 모든 인코딩 실패 시 무시하고 디코딩
                    text = file_content.decode('utf-8', errors='ignore')
            
            if not text.strip():
                raise RuntimeError("TXT 파일이 비어있습니다")
            
            return text
            
        except Exception as e:
            raise RuntimeError(f"TXT 처리 오류: {str(e)}")
